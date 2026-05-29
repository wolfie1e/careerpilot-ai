"""Resume parsing: PDF, DOCX, TXT → raw text → structured sections."""
import re
from pathlib import Path

from app.utils.text_utils import clean_text, extract_emails, extract_phones, extract_urls


def parse_file(file_path: str) -> str:
    """Extract raw text from file based on extension."""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(file_path)
    elif ext == ".docx":
        return _parse_docx(file_path)
    elif ext == ".txt":
        return _parse_txt(file_path)
    raise ValueError(f"Unsupported file type: {ext}")


def _parse_pdf(path: str) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n\n".join(pages)
        if len(text.strip()) > 100:
            return clean_text(text)
    except Exception:
        pass

    # Fallback to pdfminer
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        text = pdfminer_extract(path) or ""
        return clean_text(text)
    except Exception as e:
        raise ValueError(f"Could not parse PDF: {e}")


def _parse_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return clean_text("\n".join(paragraphs))


def _parse_txt(path: str) -> str:
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            with open(path, "r", encoding=encoding) as f:
                return clean_text(f.read())
        except UnicodeDecodeError:
            continue
    raise ValueError("Cannot decode text file")


# ── Section extraction ────────────────────────────────────────────────────────

_SECTION_PATTERNS = {
    "contact": re.compile(
        r"(contact|personal\s+info|profile)", re.IGNORECASE
    ),
    "summary": re.compile(
        r"(summary|objective|profile|about\s+me|professional\s+summary|career\s+objective)", re.IGNORECASE
    ),
    "experience": re.compile(
        r"(experience|work\s+history|employment|career\s+history|professional\s+experience)", re.IGNORECASE
    ),
    "education": re.compile(
        r"(education|academic|qualifications|degree|university|college)", re.IGNORECASE
    ),
    "skills": re.compile(
        r"(skills|technical\s+skills|core\s+competencies|technologies|tools)", re.IGNORECASE
    ),
    "projects": re.compile(
        r"(projects|personal\s+projects|portfolio|side\s+projects|open\s+source)", re.IGNORECASE
    ),
    "certifications": re.compile(
        r"(certifications?|certificates?|licenses?|credentials?)", re.IGNORECASE
    ),
    "achievements": re.compile(
        r"(achievements?|awards?|honors?|accomplishments?|publications?)", re.IGNORECASE
    ),
}


def extract_sections(raw_text: str) -> dict:
    """Split resume text into labelled sections."""
    lines = raw_text.split("\n")

    # Detect section header lines
    section_starts: list[tuple[int, str]] = []
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped or len(stripped) > 80:
            continue
        for section_name, pattern in _SECTION_PATTERNS.items():
            if pattern.search(stripped):
                section_starts.append((i, section_name))
                break

    sections: dict[str, str] = {}

    if not section_starts:
        sections["raw"] = raw_text
        contact = _extract_contact_block(lines[:10])
        if contact:
            sections["contact"] = contact
        return sections

    # Text before first section header = contact block
    first_idx = section_starts[0][0]
    if first_idx > 0:
        contact_text = "\n".join(lines[:first_idx]).strip()
        if contact_text:
            sections["contact"] = contact_text

    for idx, (line_no, sec_name) in enumerate(section_starts):
        next_line = section_starts[idx + 1][0] if idx + 1 < len(section_starts) else len(lines)
        content = "\n".join(lines[line_no + 1: next_line]).strip()
        if content:
            if sec_name in sections:
                sections[sec_name] += "\n" + content
            else:
                sections[sec_name] = content

    return sections


def _extract_contact_block(lines: list[str]) -> str:
    """Grab first non-empty lines as contact info."""
    return "\n".join(l.strip() for l in lines if l.strip())


def parse_contact_info(raw_text: str) -> dict:
    """Extract structured contact fields from resume text."""
    emails = extract_emails(raw_text)
    phones = extract_phones(raw_text)
    urls = extract_urls(raw_text)

    linkedin = next((u for u in urls if "linkedin" in u.lower()), None)
    github = next((u for u in urls if "github" in u.lower()), None)
    portfolio = next((u for u in urls if "linkedin" not in u.lower() and "github" not in u.lower()), None)

    # Name: first non-empty short line that isn't a known label
    lines = [l.strip() for l in raw_text.split("\n") if l.strip()]
    name = None
    for line in lines[:5]:
        if len(line) < 60 and not re.search(r"@|http|phone|email|address|\|", line, re.IGNORECASE):
            name = line
            break

    return {
        "name": name,
        "email": emails[0] if emails else None,
        "phone": phones[0] if phones else None,
        "linkedin": linkedin,
        "github": github,
        "portfolio": portfolio,
    }
